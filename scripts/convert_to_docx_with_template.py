#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Converte Markdown para DOCX preservando template e inserindo conteúdo após "Sumário"
Com suporte a diagramas Mermaid renderizados e syntax highlighting
Salva no diretório output/

Uso:
    python convert_to_docx_with_template.py <arquivo.md> [output.docx]
    python convert_to_docx_with_template.py --all

Dependências do sistema:
    - pandoc (brew/apt/choco install pandoc)
    - mermaid-cli (npm install -g @mermaid-js/mermaid-cli)

Dependências Python:
    - pypandoc (pip install pypandoc)
    - python-docx (pip install python-docx)
"""

import re
import os
import sys
import argparse
import subprocess
import hashlib
import tempfile
import shutil
from pathlib import Path

try:
    import pypandoc
except ImportError:
    print("[ERRO] pypandoc nao esta instalado.")
    print("   Instale com: pip install pypandoc")
    print("   Ou com Poetry: poetry add pypandoc")
    sys.exit(2)

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("[ERRO] python-docx nao esta instalado.")
    print("   Instale com: pip install python-docx")
    print("   Ou com Poetry: poetry add python-docx")
    sys.exit(2)


# Constantes
MERMAID_REGEX = r'```mermaid\n(.*?)```'


def check_dependencies():
    """Verifica se dependências do sistema estão instaladas"""
    missing = []

    if not shutil.which('mmdc'):
        missing.append('mermaid-cli (npm install -g @mermaid-js/mermaid-cli)')

    if not shutil.which('pandoc'):
        missing.append('pandoc (brew/apt/choco install pandoc)')

    if missing:
        print("[ERRO] Dependencias do sistema faltando:")
        for dep in missing:
            print(f"   - {dep}")
        sys.exit(2)


class MermaidRenderer:
    """Renderiza diagramas Mermaid para PNG usando mermaid-cli"""

    def __init__(self, cache_dir: Path):
        self.cache_dir = cache_dir
        self.cache_dir.mkdir(parents=True, exist_ok=True)

    def render(self, mermaid_code: str, index: int) -> Path:
        """Renderiza um diagrama Mermaid e retorna o caminho da imagem"""
        # Hash do código para cache
        code_hash = hashlib.md5(mermaid_code.encode()).hexdigest()[:8]
        png_path = self.cache_dir / f"diagram-{code_hash}.png"

        # Se já existe no cache, retorna
        if png_path.exists():
            print(f"  [OK] Usando cache: {png_path.name}")
            return png_path

        # Cria arquivo temporário .mmd
        with tempfile.NamedTemporaryFile(
            mode='w', suffix='.mmd', delete=False, encoding='utf-8'
        ) as tmp:
            tmp.write(mermaid_code)
            tmp_path = Path(tmp.name)

        try:
            # Renderiza com mmdc
            result = subprocess.run(
                ['mmdc', '-i', str(tmp_path), '-o', str(png_path)],
                capture_output=True,
                text=True,
                timeout=30,
                shell=True  # Necessário no Windows para encontrar mmdc.cmd
            )

            if result.returncode != 0:
                raise RuntimeError(
                    f"Erro ao renderizar Mermaid: {result.stderr}"
                )

            print(f"  [OK] Renderizado: {png_path.name}")
            return png_path

        except subprocess.TimeoutExpired:
            print(f"  [AVISO] Timeout ao renderizar diagrama {index} (>30s)")
            return None
        except Exception as e:
            print(f"  [AVISO] Erro no diagrama {index}: {e}")
            return None
        finally:
            tmp_path.unlink(missing_ok=True)


class MarkdownProcessor:
    """Processa Markdown extraindo e substituindo blocos Mermaid"""

    def __init__(self, renderer: MermaidRenderer):
        self.renderer = renderer

    def process(self, md_content: str, img_relative_path: str) -> str:
        """Substitui blocos Mermaid por imagens"""
        diagram_count = 0

        def replace_mermaid(match):
            nonlocal diagram_count
            diagram_count += 1
            mermaid_code = match.group(1)

            print(f"  Processando diagrama {diagram_count}...")
            png_path = self.renderer.render(mermaid_code, diagram_count)

            if png_path is None:
                # Retorna placeholder se falhar
                return f"![Diagrama {diagram_count} - Erro ao renderizar]"

            # Path relativo da imagem no Markdown
            # Adiciona atributo de largura para caber na página (6in = ~15cm)
            img_md_path = f"{img_relative_path}/{png_path.name}"
            return f"![Diagrama {diagram_count}]({img_md_path}){{width=6in}}"

        processed = re.sub(
            MERMAID_REGEX,
            replace_mermaid,
            md_content,
            flags=re.DOTALL
        )

        if diagram_count > 0:
            print(f"  Total: {diagram_count} diagramas processados")

        return processed


class DOCXCustomizer:
    """Customiza DOCX após conversão (mesma lógica do convert-md-to-docx.py)"""

    def _set_paragraph_shading(self, paragraph, fill_color='F5F5F5'):
        """Define cor de fundo (shading) de um parágrafo"""
        pPr = paragraph._element.get_or_add_pPr()
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), fill_color)
        shading.set(qn('w:val'), 'clear')
        pPr.append(shading)

    def _set_cell_border(self, cell, **kwargs):
        """Define bordas de uma célula de tabela"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # Extrair parâmetros comuns
        border_size = kwargs.pop('size', '4')  # Default: 0.5pt
        border_color = kwargs.pop('color', '000000')  # Default: preto

        # Criar elemento de bordas
        tcBorders = OxmlElement('w:tcBorders')

        for border_name in ('top', 'left', 'bottom', 'right'):
            if border_name in kwargs:
                border_value = kwargs[border_name]
                if border_value:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), border_value)
                    border.set(qn('w:sz'), str(border_size))
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), border_color)
                    tcBorders.append(border)

        tcPr.append(tcBorders)

    def customize_tables(self, doc):
        """Customiza tabelas: cabeçalho com borda preta, linhas com borda cinza, largura 100%"""
        for table in doc.tables:
            # Configurar tabela para ocupar 100% da largura disponível
            table.autofit = False
            table.allow_autofit = False

            # Configurar largura da tabela para 100%
            tbl = table._element
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)

            # Remover configuração de largura fixa se existir
            tblW = tblPr.find(qn('w:tblW'))
            if tblW is not None:
                tblPr.remove(tblW)

            # Definir largura como 100% (5000 = 100% em unidades do Word)
            tblW = OxmlElement('w:tblW')
            tblW.set(qn('w:w'), '5000')
            tblW.set(qn('w:type'), 'pct')  # Percentual
            tblPr.append(tblW)

            # Customizar bordas das células
            for row_idx, row in enumerate(table.rows):
                # Primeira linha = cabeçalho (borda preta grossa)
                is_header = (row_idx == 0)

                for cell in row.cells:
                    if is_header:
                        # Cabeçalho: borda preta grossa
                        self._set_cell_border(
                            cell,
                            top='single',
                            bottom='single',
                            left='nil',
                            right='nil',
                            size='16',      # Grossa (2pt)
                            color='000000'  # Preta
                        )
                    else:
                        # Linhas normais: borda cinza fina
                        self._set_cell_border(
                            cell,
                            top='single',
                            bottom='single',
                            left='nil',
                            right='nil',
                            size='4',       # Fina (0.5pt)
                            color='CCCCCC'  # Cinza
                        )

    def customize_code_blocks(self, doc):
        """Adiciona fundo cinza claro aos blocos de código com ``` no MD"""
        code_block_styles = ['Source Code', 'Verbatim', 'Code',
                             'Preformatted Text', 'HTML Preformatted']

        for paragraph in doc.paragraphs:
            if paragraph.style and paragraph.style.name in code_block_styles:
                # Verificar se NÃO é lista (numPr)
                pPr = paragraph._element.pPr
                is_list = False
                if pPr is not None:
                    numPr = pPr.find(qn('w:numPr'))
                    is_list = numPr is not None

                # Apenas aplicar fundo se não for lista
                if not is_list:
                    self._set_paragraph_shading(paragraph, 'F5F5F5')  # Cinza claro

    def ensure_list_formatting(self, doc):
        """Garante que listas tenham bullets/números visíveis"""
        for paragraph in doc.paragraphs:
            # Verificar se o parágrafo tem numeração (numPr)
            pPr = paragraph._element.get_or_add_pPr()
            numPr = pPr.find(qn('w:numPr'))

            if numPr is not None:
                # O parágrafo faz parte de uma lista
                numId_elem = numPr.find(qn('w:numId'))
                ilvl_elem = numPr.find(qn('w:ilvl'))

                if numId_elem is not None and ilvl_elem is not None:
                    # Garantir que o parágrafo tenha o estilo de lista adequado
                    current_style = paragraph.style.name if paragraph.style else ''

                    if 'List' not in current_style:
                        try:
                            paragraph.style = 'List Paragraph'
                        except:
                            pass

    def set_margins(self, doc, left=0.5, right=0.5, top=0.5, bottom=0.5):
        """Define margens do documento em polegadas"""
        sections = doc.sections
        for section in sections:
            section.left_margin = Inches(left)
            section.right_margin = Inches(right)
            section.top_margin = Inches(top)
            section.bottom_margin = Inches(bottom)
        print(f"  [OK] Margens configuradas: {left}in (esq/dir), {top}in (sup/inf)")

    def remove_all_bookmarks(self, doc):
        """Remove todos os bookmarks do documento (criados pelo Pandoc)"""
        bookmark_count = 0

        # Os bookmarks estão no body do documento
        body = doc.element.body

        # Remover todos os bookmarkStart e bookmarkEnd do body
        for child in list(body):
            if child.tag == qn('w:bookmarkStart'):
                # Não remover bookmarks especiais (_GoBack, etc)
                name = child.get(qn('w:name'))
                if name and not name.startswith('_'):
                    body.remove(child)
                    bookmark_count += 1
            elif child.tag == qn('w:bookmarkEnd'):
                body.remove(child)

        # Também remover bookmarks dentro de parágrafos
        for paragraph in doc.paragraphs:
            for element in list(paragraph._element):
                if element.tag == qn('w:bookmarkStart'):
                    name = element.get(qn('w:name'))
                    if name and not name.startswith('_'):
                        paragraph._element.remove(element)
                        bookmark_count += 1
                elif element.tag == qn('w:bookmarkEnd'):
                    paragraph._element.remove(element)

        if bookmark_count > 0:
            print(f"  [OK] Removidos {bookmark_count} bookmarks")

    def remove_hyperlinks_from_headings(self, doc):
        """Remove hyperlinks dos títulos (headings)"""
        removed_count = 0

        for paragraph in doc.paragraphs:
            # Verificar se é um heading
            if paragraph.style and paragraph.style.name.startswith('Heading'):
                # Procurar hyperlinks
                for element in list(paragraph._element):
                    if element.tag == qn('w:hyperlink'):
                        # Extrair o texto dos runs dentro do hyperlink
                        text_runs = []
                        for run_elem in element.findall(qn('w:r')):
                            text_runs.append(run_elem)

                        # Remover o hyperlink e inserir os runs diretamente
                        parent = element.getparent()
                        index = list(parent).index(element)
                        parent.remove(element)

                        # Reinserir os runs na mesma posição
                        for run_elem in reversed(text_runs):
                            parent.insert(index, run_elem)

                        removed_count += 1

        if removed_count > 0:
            print(f"  [OK] Removidos {removed_count} hyperlinks de titulos")

    def fix_toc(self, doc):
        """Substitui 'Table of Contents' por 'Sumário'"""
        # Procurar e substituir "Table of Contents" por "Sumário"
        for para in doc.paragraphs:
            text = para.text.strip()
            if text in ['Table of Contents', 'TABLE OF CONTENTS']:
                para.text = 'Sumário'
                print(f"  [OK] Substituido '{text}' por 'Sumario'")
                break
            elif text in ['Sumário', 'SUMÁRIO']:
                print(f"  [OK] Titulo 'Sumario' ja existe")
                break

    def customize_styles(self, doc):
        """Customiza estilos de fonte"""
        styles = doc.styles

        # Estilo Normal (texto padrão)
        if 'Normal' in styles:
            style = styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(11)

        # Estilo Title (título do documento) - Arial
        if 'Title' in styles:
            style = styles['Title']
            style.font.name = 'Arial'
            style.font.size = Pt(18)

        # Estilos de Heading (títulos) - todos em Arial
        for i in range(1, 10):
            heading_name = f'Heading {i}'
            if heading_name in styles:
                style = styles[heading_name]
                style.font.name = 'Arial'
                # Tamanhos decrescentes para headings
                sizes = {1: 16, 2: 14, 3: 13, 4: 12, 5: 11, 6: 11, 7: 11, 8: 11, 9: 11}
                style.font.size = Pt(sizes.get(i, 11))

        # Estilos de código inline (entre ` `) - Courier New
        inline_code_styles = ['Verbatim Char']
        for style_name in inline_code_styles:
            if style_name in styles:
                style = styles[style_name]
                style.font.name = 'Courier New'
                style.font.size = Pt(11)

        # Estilos de blocos de código - Courier New
        block_code_styles = ['Source Code', 'Verbatim', 'Code',
                             'Preformatted Text', 'HTML Preformatted']
        for style_name in block_code_styles:
            if style_name in styles:
                style = styles[style_name]
                style.font.name = 'Courier New'
                style.font.size = Pt(11)

        # Estilos de lista - Arial
        list_styles = ['Compact', 'List Paragraph', 'List', 'List Bullet', 'List Number']
        for style_name in list_styles:
            if style_name in styles:
                style = styles[style_name]
                style.font.name = 'Arial'
                style.font.size = Pt(11)

    def customize_docx(self, doc):
        """Aplica todas as customizações (mesma lógica do convert-md-to-docx.py)"""
        # Configurar margens
        self.set_margins(doc, left=0.5, right=0.5, top=0.5, bottom=0.5)

        # Customizar estilos de fonte
        self.customize_styles(doc)

        # Customizar tabelas (apenas bordas horizontais)
        self.customize_tables(doc)

        # Garantir formatação de listas (bullets e números)
        self.ensure_list_formatting(doc)

        # Adicionar fundo cinza APENAS aos blocos de código (não listas)
        self.customize_code_blocks(doc)

        # Corrigir TOC
        self.fix_toc(doc)

        # Remover todos os bookmarks
        self.remove_all_bookmarks(doc)

        # Remover hyperlinks dos títulos
        self.remove_hyperlinks_from_headings(doc)

        print(f"  [OK] Customizacoes aplicadas (estilos + tabelas + codigo + listas + TOC)")


def convert_md_to_docx_with_pandoc(md_content: str, output_path: Path):
    """Converte MD para DOCX usando Pandoc (com syntax highlighting)"""
    # Cria arquivo temporário para MD processado
    with tempfile.NamedTemporaryFile(
        mode='w',
        suffix='.md',
        delete=False,
        encoding='utf-8'
    ) as tmp:
        tmp.write(md_content)
        tmp_path = tmp.name

    try:
        # Caminho do arquivo de referência (se existir)
        script_dir = Path(__file__).parent
        reference_docx = script_dir / 'templates' / 'reference.docx'

        # Metadata para Pandoc
        extra_args = [
            '--standalone',
            '--wrap=auto',
            '--highlight-style=pygments',  # Syntax highlighting com cores
        ]

        # Usar arquivo de referência se existir
        if reference_docx.exists():
            extra_args.append(f'--reference-doc={reference_docx}')

        # Conversão via pypandoc
        pypandoc.convert_file(
            tmp_path,
            'docx',
            outputfile=str(output_path),
            extra_args=extra_args
        )

        print(f"  [OK] Conversao MD -> DOCX via Pandoc concluida")

    finally:
        Path(tmp_path).unlink(missing_ok=True)


def copy_header_footer(source_doc, target_doc):
    """Copia header e footer do documento source para o target"""
    from copy import deepcopy

    for src_section, tgt_section in zip(source_doc.sections, target_doc.sections):
        # Copiar header
        src_header = src_section.header
        tgt_header = tgt_section.header

        # Limpar header atual mantendo estrutura
        for element in list(tgt_header._element):
            # Remover apenas parágrafos e tabelas, manter outros elementos (settings, etc.)
            if element.tag.endswith('}p') or element.tag.endswith('}tbl'):
                tgt_header._element.remove(element)

        # Copiar apenas parágrafos e tabelas do header source
        for element in src_header._element:
            if element.tag.endswith('}p') or element.tag.endswith('}tbl'):
                new_element = deepcopy(element)
                tgt_header._element.append(new_element)

        # Copiar footer
        src_footer = src_section.footer
        tgt_footer = tgt_section.footer

        # Limpar footer atual mantendo estrutura
        for element in list(tgt_footer._element):
            if element.tag.endswith('}p') or element.tag.endswith('}tbl'):
                tgt_footer._element.remove(element)

        # Copiar apenas parágrafos e tabelas do footer source
        for element in src_footer._element:
            if element.tag.endswith('}p') or element.tag.endswith('}tbl'):
                new_element = deepcopy(element)
                tgt_footer._element.append(new_element)

    print(f"  [OK] Header e footer do template copiados")


def merge_template_with_content(template_path: str, content_docx_path: Path, output_path: str):
    """Mescla template com conteúdo do DOCX gerado pelo Pandoc"""
    print(f"  [OK] Mesclando template com conteudo...")

    # Carrega template (preserva conteúdo)
    doc_template = Document(template_path)

    # Carrega DOCX gerado pelo Pandoc
    doc_content = Document(str(content_docx_path))

    # Copia header e footer do template para o documento de conteúdo
    # (fazemos isso porque vamos usar o doc_content como base, que tem melhor formatação)
    copy_header_footer(doc_template, doc_content)

    # Copia elementos do template para o INÍCIO do documento de conteúdo
    print(f"  [OK] Adicionando {len(doc_template.element.body)} elementos do template no inicio...")

    # Inserir elementos do template no início do body do content
    content_body = doc_content.element.body
    for i, element in enumerate(doc_template.element.body):
        content_body.insert(i, element)

    # Aplica customizações (mesma lógica do convert-md-to-docx.py)
    print(f"  [OK] Aplicando customizacoes...")
    customizer = DOCXCustomizer()
    customizer.customize_docx(doc_content)

    # Salva o resultado
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc_content.save(output_path)

    print(f"  [OK] Mesclagem concluida")


def convert_md_to_docx(input_path, output_path, template_path, cache_dir):
    """Converte um arquivo MD para DOCX usando template"""
    print(f"\n{'='*60}")
    print(f"Convertendo: {Path(input_path).name}")
    print(f"{'='*60}")
    print(f"Entrada: {input_path}")
    print(f"Template: {template_path}")
    print(f"Saida: {output_path}")

    # Lê MD
    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Remove comentários HTML
    content = re.sub(r'<!--.*?-->', '', content, flags=re.DOTALL)

    # Processa Mermaid
    renderer = MermaidRenderer(cache_dir)
    processor = MarkdownProcessor(renderer)

    # Path ABSOLUTO das imagens (para que Pandoc encontre de qualquer lugar)
    img_path = str(cache_dir.absolute()).replace('\\', '/')
    processed_md = processor.process(content, img_path)

    # Cria DOCX temporário com Pandoc (com syntax highlighting)
    temp_docx = Path(tempfile.mktemp(suffix='.docx'))

    try:
        print(f"  [OK] Convertendo MD -> DOCX com Pandoc...")
        convert_md_to_docx_with_pandoc(processed_md, temp_docx)

        # Mescla template com conteúdo
        merge_template_with_content(template_path, temp_docx, output_path)

        print(f"  [OK] Arquivo salvo: {output_path}")
        print(f"{'='*60}\n")

    finally:
        # Remove arquivo temporário
        if temp_docx.exists():
            temp_docx.unlink()


def main():
    parser = argparse.ArgumentParser(
        description='Converte Markdown para DOCX usando template',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Exemplos:
  %(prog)s arquivo.md
  %(prog)s arquivo.md saida.docx
  %(prog)s --all
        """
    )
    parser.add_argument('input', nargs='?', help='Arquivo .md de entrada ou --all')
    parser.add_argument('output', nargs='?', help='Arquivo .docx de saída (opcional)')
    parser.add_argument('--all', action='store_true', help='Converter todos os .md em docs/relatorios/')

    args = parser.parse_args()

    # Verifica dependências
    check_dependencies()

    # Caminhos padrão
    template_path = 'docs/relatorios/templates/Template Relatório-Técnico-DestaquesGovbr Tema 7.docx'
    relatorios_dir = 'docs/relatorios'
    output_dir = 'docs/relatorios/output'
    cache_dir = Path(output_dir) / 'imgs'

    # Valida template
    if not os.path.exists(template_path):
        print(f"[ERRO] Template nao encontrado: {template_path}")
        sys.exit(1)

    # Modo --all
    if args.all or (args.input == '--all'):
        md_files = list(Path(relatorios_dir).glob('*.md'))
        if not md_files:
            print(f"[ERRO] Nenhum arquivo .md encontrado em {relatorios_dir}")
            sys.exit(1)

        print(f"\n{'='*60}")
        print(f"Convertendo {len(md_files)} arquivo(s)...")
        print(f"{'='*60}")

        for md_file in md_files:
            input_path = str(md_file)
            output_filename = md_file.stem + '.docx'
            output_path = os.path.join(output_dir, output_filename)

            try:
                convert_md_to_docx(input_path, output_path, template_path, cache_dir)
            except Exception as e:
                print(f"[ERRO] Erro ao converter {md_file.name}: {e}")
                import traceback
                traceback.print_exc()

        print(f"\n{'='*60}")
        print(f"[OK] Conversao concluida!")
        print(f"{'='*60}")
        return

    # Modo arquivo único
    if not args.input:
        parser.print_help()
        sys.exit(1)

    input_path = args.input

    # Ajusta caminho se for apenas nome do arquivo
    if not os.path.exists(input_path):
        input_path = os.path.join(relatorios_dir, args.input)

    if not os.path.exists(input_path):
        print(f"[ERRO] Arquivo nao encontrado: {args.input}")
        sys.exit(1)

    # Define output
    if args.output:
        output_path = args.output
    else:
        input_stem = Path(input_path).stem
        output_path = os.path.join(output_dir, f"{input_stem}.docx")

    try:
        convert_md_to_docx(input_path, output_path, template_path, cache_dir)
        print(f"[OK] Conversao concluida com sucesso!")
    except Exception as e:
        print(f"[ERRO] Erro: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()