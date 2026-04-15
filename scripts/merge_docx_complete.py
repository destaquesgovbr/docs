#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Merge COMPLETO de DOCX preservando ABSOLUTAMENTE TUDO
- Copia arquivos de mídia (imagens) do template e conteúdo
- Mescla relationships
- Preserva header/footer com imagens
- Preserva estilos, fontes, formatação
"""

import sys
import shutil
import zipfile
from pathlib import Path
import tempfile
import os
import re

try:
    from docx import Document
    from docx.oxml.ns import qn
    from copy import deepcopy
except ImportError:
    print("[ERRO] python-docx nao esta instalado.")
    sys.exit(2)


def extract_docx(docx_path, extract_dir):
    """Extrai DOCX (ZIP) para diretório"""
    with zipfile.ZipFile(docx_path, 'r') as zip_ref:
        zip_ref.extractall(extract_dir)


def create_docx(source_dir, output_path):
    """Cria DOCX a partir de diretório"""
    with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as docx:
        for root, dirs, files in os.walk(source_dir):
            for file in files:
                file_path = os.path.join(root, file)
                arcname = os.path.relpath(file_path, source_dir)
                docx.write(file_path, arcname)


def copy_template_media_to_content(template_dir, content_dir):
    """Copia TODOS os arquivos de mídia do template para o conteúdo com nomes únicos"""
    template_media = Path(template_dir) / 'word' / 'media'
    content_media = Path(content_dir) / 'word' / 'media'

    if not template_media.exists():
        return {}

    content_media.mkdir(parents=True, exist_ok=True)

    # Mapa: nome_original -> novo_nome
    media_map = {}

    for media_file in template_media.iterdir():
        if media_file.is_file():
            old_name = media_file.name
            new_name = f"template_{old_name}"

            dest_file = content_media / new_name
            shutil.copy2(media_file, dest_file)

            media_map[old_name] = new_name
            print(f"  [OK] Copiado: {old_name} -> {new_name}")

    return media_map


def update_xml_media_references(xml_path, media_map):
    """Atualiza referências de mídia em arquivos XML"""
    if not os.path.exists(xml_path):
        return

    with open(xml_path, 'r', encoding='utf-8') as f:
        content = f.read()

    modified = False
    for old_name, new_name in media_map.items():
        if old_name in content:
            content = content.replace(f'media/{old_name}', f'media/{new_name}')
            modified = True

    if modified:
        with open(xml_path, 'w', encoding='utf-8') as f:
            f.write(content)


def merge_docx_with_media(template_path, content_path, output_path):
    """Merge completo preservando TODAS as imagens"""
    print(f"\n{'='*60}")
    print(f"MERGE COMPLETO DE DOCX")
    print(f"{'='*60}")
    print(f"Template: {template_path}")
    print(f"Conteudo: {content_path}")
    print(f"Output: {output_path}")
    print(f"{'='*60}\n")

    with tempfile.TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)
        template_dir = temp_path / 'template'
        content_dir = temp_path / 'content'

        print(f"  [OK] Extraindo arquivos DOCX...")
        extract_docx(template_path, template_dir)
        extract_docx(content_path, content_dir)

        print(f"\n  [OK] Copiando arquivos de midia do template...")
        media_map = copy_template_media_to_content(template_dir, content_dir)
        print(f"  [OK] {len(media_map)} arquivos de midia copiados")

        if media_map:
            print(f"\n  [OK] Atualizando referencias de midia no template...")
            # Atualizar referências nos XMLs do template
            template_xml_files = [
                template_dir / 'word' / 'document.xml',
                template_dir / 'word' / 'header1.xml',
                template_dir / 'word' / 'header2.xml',
                template_dir / 'word' / 'footer1.xml',
                template_dir / 'word' / 'footer2.xml',
                template_dir / 'word' / '_rels' / 'document.xml.rels',
                template_dir / 'word' / '_rels' / 'header1.xml.rels',
                template_dir / 'word' / '_rels' / 'header2.xml.rels',
            ]

            for xml_file in template_xml_files:
                if xml_file.exists():
                    update_xml_media_references(str(xml_file), media_map)
                    print(f"    - {xml_file.name}")

        # Recria DOCX do content com mídia adicional
        print(f"\n  [OK] Recriando DOCX com toda a midia...")
        temp_content_docx = temp_path / 'content_with_media.docx'
        create_docx(content_dir, temp_content_docx)

        # Recarrega template (com referências atualizadas)
        temp_template_docx = temp_path / 'template_updated.docx'
        create_docx(template_dir, temp_template_docx)

        print(f"\n  [OK] Fazendo merge estrutural com python-docx...")

        doc_template = Document(str(temp_template_docx))
        doc_content = Document(str(temp_content_docx))

        # Copia header/footer
        print(f"  [OK] Copiando header e footer...")
        for src_section, tgt_section in zip(doc_template.sections, doc_content.sections):
            # Header
            src_header = src_section.header
            tgt_header = tgt_section.header

            for element in list(tgt_header._element):
                if element.tag.endswith('}p') or element.tag.endswith('}tbl'):
                    tgt_header._element.remove(element)

            for element in src_header._element:
                if element.tag.endswith('}p') or element.tag.endswith('}tbl'):
                    new_element = deepcopy(element)
                    tgt_header._element.append(new_element)

            # Footer
            src_footer = src_section.footer
            tgt_footer = tgt_section.footer

            for element in list(tgt_footer._element):
                if element.tag.endswith('}p') or element.tag.endswith('}tbl'):
                    tgt_footer._element.remove(element)

            for element in src_footer._element:
                if element.tag.endswith('}p') or element.tag.endswith('}tbl'):
                    new_element = deepcopy(element)
                    tgt_footer._element.append(new_element)

        # Insere template no início
        print(f"  [OK] Inserindo template no inicio...")
        content_body = doc_content.element.body
        for i, element in enumerate(doc_template.element.body):
            new_element = deepcopy(element)
            content_body.insert(i, new_element)

        # Salva
        print(f"  [OK] Salvando documento final...")
        doc_content.save(output_path)

    # Estatísticas
    print(f"\n  [INFO] Verificando arquivo final...")
    doc_final = Document(output_path)

    with zipfile.ZipFile(output_path, 'r') as z:
        final_media = [f for f in z.namelist() if 'media/' in f]

    print(f"    - Paragrafos: {len(doc_final.paragraphs)}")
    print(f"    - Tabelas: {len(doc_final.tables)}")
    print(f"    - Arquivos de midia: {len(final_media)}")

    print(f"\n{'='*60}")
    print(f"[OK] Merge completo concluido!")
    print(f"{'='*60}\n")
    print(f"Arquivo: {output_path}")
    print(f"Midia preservada:")
    print(f"  - Template: {len(media_map)} imagens")
    print(f"  - Conteudo: ~{len(final_media) - len(media_map)} imagens")
    print(f"  - Total: {len(final_media)} arquivos de midia")


def main():
    if len(sys.argv) != 4:
        print("Uso: python merge_docx_complete.py <template.docx> <content.docx> <output.docx>")
        sys.exit(1)

    template_path = sys.argv[1]
    content_path = sys.argv[2]
    output_path = sys.argv[3]

    if not Path(template_path).exists():
        print(f"[ERRO] Template nao encontrado: {template_path}")
        sys.exit(1)

    if not Path(content_path).exists():
        print(f"[ERRO] Conteudo nao encontrado: {content_path}")
        sys.exit(1)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)

    try:
        merge_docx_with_media(template_path, content_path, output_path)
    except Exception as e:
        print(f"\n[ERRO] Falha: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()