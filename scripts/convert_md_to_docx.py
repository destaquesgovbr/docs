#!/usr/bin/env python3
"""
Script para converter relatórios Markdown (.md) para DOCX (.docx)
com renderização de diagramas Mermaid para imagens PNG.

Uso:
    python convert_md_to_docx.py <input.md> [--output output.docx]
    python convert_md_to_docx.py --all  # Converte todos .md em relatorios/

Dependências do sistema:
    - pandoc (brew install pandoc / apt install pandoc)
    - mermaid-cli (npm install -g @mermaid-js/mermaid-cli)

Dependências Python:
    - pypandoc (pip install pypandoc)
    - python-docx (pip install python-docx)
"""

import argparse
import re
import subprocess
import hashlib
import tempfile
import shutil
import sys
import zipfile
from pathlib import Path
from typing import List, Tuple

try:
    import pypandoc
except ImportError:
    print("[X] Erro: pypandoc nao esta instalado.")
    print("   Instale com: pip install pypandoc")
    print("   Ou com Poetry: poetry add pypandoc")
    sys.exit(2)

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("[X] Erro: python-docx nao esta instalado.")
    print("   Instale com: pip install python-docx")
    print("   Ou com Poetry: poetry add python-docx")
    sys.exit(2)


# Constantes
MERMAID_REGEX = r'```mermaid\n(.*?)```'
IMG_DIR = 'imgs'
OUTPUT_DIR = 'output'


def check_dependencies():
    """Verifica se dependências do sistema estão instaladas"""
    missing = []

    if not shutil.which('mmdc'):
        missing.append('mermaid-cli (npm install -g @mermaid-js/mermaid-cli)')

    if not shutil.which('pandoc'):
        missing.append('pandoc (brew/apt/choco install pandoc)')

    if missing:
        print("[X] Dependencias do sistema faltando:")
        for dep in missing:
            print(f"   - {dep}")
        print("\nConsulte scripts/README-convert.md para instrucoes de instalacao.")
        sys.exit(2)


class MermaidRenderer:
    """Renderiza diagramas Mermaid para PNG usando mermaid-cli"""

    def __init__(self, cache_dir: Path):
        self.cache_dir = cache_dir
        self.cache_dir.mkdir(parents=True, exist_ok=True)

    def render(self, mermaid_code: str, index: int) -> Path:
        """Renderiza um diagrama Mermaid e retorna o caminho da imagem"""
        # Hash do código para cache
        code_hash = hashlib.md5(mermaid_code.encode()).hexdigest()[:8]
        png_path = self.cache_dir / f"diagram-{code_hash}.png"

        # Se já existe no cache, retorna
        if png_path.exists():
            print(f"  [OK] Usando cache: {png_path.name}")
            return png_path

        # Cria arquivo temporário .mmd
        with tempfile.NamedTemporaryFile(
            mode='w', suffix='.mmd', delete=False, encoding='utf-8'
        ) as tmp:
            tmp.write(mermaid_code)
            tmp_path = Path(tmp.name)

        try:
            # Renderiza com mmdc
            result = subprocess.run(
                ['mmdc', '-i', str(tmp_path), '-o', str(png_path)],
                capture_output=True,
                text=True,
                timeout=30,
                shell=True  # Necessário no Windows para encontrar mmdc.cmd
            )

            if result.returncode != 0:
                raise RuntimeError(
                    f"Erro ao renderizar Mermaid: {result.stderr}"
                )

            print(f"  [OK] Renderizado: {png_path.name}")
            return png_path

        except subprocess.TimeoutExpired:
            print(f"  [!]  Timeout ao renderizar diagrama {index} (>30s)")
            return self.create_placeholder(f"Timeout no diagrama {index}")
        except Exception as e:
            print(f"  [!]  Erro no diagrama {index}: {e}")
            return self.create_placeholder(f"Erro ao renderizar diagrama {index}")
        finally:
            tmp_path.unlink(missing_ok=True)

    def create_placeholder(self, message: str) -> Path:
        """Cria uma imagem placeholder para diagramas com erro"""
        placeholder_path = self.cache_dir / f"placeholder-{hashlib.md5(message.encode()).hexdigest()[:8]}.txt"
        placeholder_path.write_text(message, encoding='utf-8')
        print(f"  [!]  Placeholder criado: {placeholder_path.name}")
        return placeholder_path


class MarkdownProcessor:
    """Processa Markdown extraindo e substituindo blocos Mermaid"""

    def __init__(self, renderer: MermaidRenderer):
        self.renderer = renderer

    def remove_emojis(self, text: str) -> str:
        """Remove emojis e ícones do texto, mas preserva box-drawing characters e status icons"""
        # Padrão que captura a maioria dos emojis Unicode
        # NOTA: Quebrado em ranges para EXCLUIR:
        #   - U+2500-U+257F (box-drawing characters)
        #   - U+2705 (✅ WHITE HEAVY CHECK MARK - usado em tabelas de status)
        #   - U+274C (❌ CROSS MARK - usado em tabelas de status)
        emoji_pattern = re.compile(
            "["
            "\U0001F600-\U0001F64F"  # emoticons
            "\U0001F300-\U0001F5FF"  # symbols & pictographs
            "\U0001F680-\U0001F6FF"  # transport & map symbols
            "\U0001F1E0-\U0001F1FF"  # flags (iOS)
            "\U00002702-\U000027B0"  # dingbats (range original)
            "\U000024C2-\U000024FF"  # enclosed characters (antes de box-drawing)
            # PULAR U+2500-U+257F (box-drawing characters - preservar!)
            "\U00002580-\U000025FF"  # block elements
            "\u2600-\u26FF"          # miscellaneous symbols
            "\u2700-\u2704"          # dingbats (antes do ✅)
            # PULAR U+2705 (✅ - preservar!)
            "\u2706-\u274B"          # dingbats (entre ✅ e ❌)
            # PULAR U+274C (❌ - preservar!)
            "\u274D-\u27BF"          # dingbats (depois do ❌)
            "\U0001F251-\U0001F251"  # enclosed characters (final)
            "\U0001F900-\U0001F9FF"  # supplemental symbols
            "\U0001FA00-\U0001FA6F"  # extended symbols
            "]+",
            flags=re.UNICODE
        )
        return emoji_pattern.sub('', text)

    def fix_lists_after_colon(self, text: str) -> str:
        """
        Corrige listas que aparecem após ':' para garantir formatação correta
        Pandoc às vezes junta essas listas em um único parágrafo
        """
        lines = text.split('\n')
        fixed_lines = []

        for i, line in enumerate(lines):
            fixed_lines.append(line)

            # Se a linha termina com ':' e a próxima é uma lista
            if line.strip().endswith(':') and i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                # Se a próxima linha é uma lista (começa com - ou número.)
                if next_line.startswith('- ') or (len(next_line) > 2 and next_line[0].isdigit() and next_line[1:3] in ['. ', ') ']):
                    # Adicionar linha em branco para forçar o Pandoc a reconhecer a lista
                    fixed_lines.append('')

        return '\n'.join(fixed_lines)

    def process(self, md_content: str, img_relative_path: str) -> str:
        """Substitui blocos Mermaid por imagens"""
        # Corrige listas após ':'
        md_content = self.fix_lists_after_colon(md_content)

        diagram_count = 0

        def replace_mermaid(match):
            nonlocal diagram_count
            diagram_count += 1
            mermaid_code = match.group(1)

            print(f"  Processando diagrama {diagram_count}...")
            png_path = self.renderer.render(mermaid_code, diagram_count)

            # Path relativo da imagem no Markdown
            # Adiciona atributo de largura para caber na página (6in = ~15cm)
            img_md_path = f"{img_relative_path}/{png_path.name}"
            return f"![Diagrama {diagram_count}]({img_md_path}){{width=6in}}"

        processed = re.sub(
            MERMAID_REGEX,
            replace_mermaid,
            md_content,
            flags=re.DOTALL
        )

        print(f"  Total: {diagram_count} diagramas processados")
        return processed


class DOCXConverter:
    """Converte Markdown para DOCX usando Pandoc"""

    def _set_paragraph_shading(self, paragraph, fill_color='F5F5F5'):
        """
        Define cor de fundo (shading) de um parágrafo
        fill_color: cor em hex (ex: 'F5F5F5' para cinza claro)
        """
        pPr = paragraph._element.get_or_add_pPr()
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), fill_color)
        shading.set(qn('w:val'), 'clear')
        pPr.append(shading)

    def _set_cell_border(self, cell, **kwargs):
        """
        Define bordas de uma célula de tabela
        kwargs:
          - top, bottom, left, right: 'single', 'nil' ou None
          - size: tamanho em oitavos de ponto (4 = 0.5pt, 16 = 2pt)
          - color: cor hex (ex: '000000', 'CCCCCC')
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # Extrair parâmetros comuns
        border_size = kwargs.pop('size', '4')  # Default: 0.5pt
        border_color = kwargs.pop('color', '000000')  # Default: preto

        # Criar elemento de bordas
        tcBorders = OxmlElement('w:tcBorders')

        for border_name in ('top', 'left', 'bottom', 'right'):
            if border_name in kwargs:
                border_value = kwargs[border_name]
                if border_value:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), border_value)
                    border.set(qn('w:sz'), str(border_size))
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), border_color)
                    tcBorders.append(border)

        tcPr.append(tcBorders)

    def customize_tables(self, doc):
        """Customiza tabelas: cabeçalho com borda preta, linhas com borda cinza, largura 100%"""
        for table in doc.tables:
            # Configurar tabela para ocupar 100% da largura disponível
            table.autofit = False
            table.allow_autofit = False

            # Configurar largura da tabela para 100%
            tbl = table._element
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)

            # Remover configuração de largura fixa se existir
            tblW = tblPr.find(qn('w:tblW'))
            if tblW is not None:
                tblPr.remove(tblW)

            # Definir largura como 100% (5000 = 100% em unidades do Word)
            tblW = OxmlElement('w:tblW')
            tblW.set(qn('w:w'), '5000')
            tblW.set(qn('w:type'), 'pct')  # Percentual
            tblPr.append(tblW)

            # Customizar bordas das células
            for row_idx, row in enumerate(table.rows):
                # Primeira linha = cabeçalho (borda preta grossa)
                is_header = (row_idx == 0)

                for cell in row.cells:
                    if is_header:
                        # Cabeçalho: borda preta grossa
                        self._set_cell_border(
                            cell,
                            top='single',
                            bottom='single',
                            left='nil',
                            right='nil',
                            size='16',      # Grossa (2pt)
                            color='000000'  # Preta
                        )
                    else:
                        # Linhas normais: borda cinza fina
                        self._set_cell_border(
                            cell,
                            top='single',
                            bottom='single',
                            left='nil',
                            right='nil',
                            size='4',       # Fina (0.5pt)
                            color='CCCCCC'  # Cinza
                        )

    def customize_code_blocks(self, doc):
        """Adiciona fundo cinza claro aos blocos de código com ``` no MD"""
        from docx.oxml.ns import qn

        # APENAS Source Code = blocos com ```
        # NÃO incluir Compact (usado erroneamente para listas pelo Pandoc)
        code_block_styles = ['Source Code', 'Verbatim', 'Code',
                             'Preformatted Text', 'HTML Preformatted']

        for paragraph in doc.paragraphs:
            # Verificar se o parágrafo usa estilo de código E não faz parte de lista
            if paragraph.style and paragraph.style.name in code_block_styles:
                # Verificar se NÃO é lista (numPr)
                pPr = paragraph._element.pPr
                is_list = False
                if pPr is not None:
                    numPr = pPr.find(qn('w:numPr'))
                    is_list = numPr is not None

                # Apenas aplicar fundo se não for lista
                if not is_list:
                    self._set_paragraph_shading(paragraph, 'F5F5F5')  # Cinza claro

    def ensure_list_formatting(self, doc):
        """
        Garante que listas tenham bullets/números visíveis
        Corrige problema onde Pandoc gera listas sem marcadores visíveis
        """
        from docx.oxml.shared import OxmlElement, qn

        for paragraph in doc.paragraphs:
            # Verificar se o parágrafo tem numeração (numPr)
            pPr = paragraph._element.get_or_add_pPr()
            numPr = pPr.find(qn('w:numPr'))

            if numPr is not None:
                # O parágrafo faz parte de uma lista
                # Garantir que o estilo seja apropriado

                # Verificar se é lista com bullet ou numerada
                numId_elem = numPr.find(qn('w:numId'))
                ilvl_elem = numPr.find(qn('w:ilvl'))

                if numId_elem is not None and ilvl_elem is not None:
                    # Garantir que o parágrafo tenha o estilo de lista adequado
                    # Se não tiver estilo de lista, aplicar List Paragraph
                    current_style = paragraph.style.name if paragraph.style else ''

                    if 'List' not in current_style:
                        try:
                            paragraph.style = 'List Paragraph'
                        except:
                            # Se não existir, criar estilo básico de lista
                            pass

    def set_margins(self, doc, left=0.5, right=0.5, top=0.5, bottom=0.5):
        """
        Define margens do documento
        Valores em polegadas (inches)
        """
        sections = doc.sections
        for section in sections:
            section.left_margin = Inches(left)
            section.right_margin = Inches(right)
            section.top_margin = Inches(top)
            section.bottom_margin = Inches(bottom)
        print(f"  [OK] Margens configuradas: {left}in (esq/dir), {top}in (sup/inf)")

    def remove_all_bookmarks(self, doc):
        """Remove todos os bookmarks do documento (criados pelo Pandoc)"""
        bookmark_count = 0

        # Os bookmarks estão no body do documento, não nos parágrafos
        body = doc.element.body

        # Remover todos os bookmarkStart e bookmarkEnd do body
        for child in list(body):
            if child.tag == qn('w:bookmarkStart'):
                # Não remover bookmarks especiais (_GoBack, etc)
                name = child.get(qn('w:name'))
                if name and not name.startswith('_'):
                    body.remove(child)
                    bookmark_count += 1
            elif child.tag == qn('w:bookmarkEnd'):
                body.remove(child)

        print(f"  [OK] Removidos {bookmark_count} bookmarks dos títulos")

    def fix_toc(self, doc):
        """Substitui 'Table of Contents' por 'Sumário' e remove bookmarks"""
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        import re

        # Procurar e substituir "Table of Contents" por "Sumário"
        toc_start_idx = None

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text in ['Table of Contents', 'TABLE OF CONTENTS']:
                # Substituir por "Sumário"
                para.text = 'Sumário'
                print(f"  [OK] Substituído '{text}' por 'Sumário'")
                toc_start_idx = i
                break
            elif text in ['Sumário', 'SUMÁRIO']:
                # Já está correto
                print(f"  [OK] Título 'Sumário' já existe")
                toc_start_idx = i
                break

        # Remover todos os bookmarks (criados automaticamente pelo Pandoc)
        self.remove_all_bookmarks(doc)

    def _add_hyperlink(self, paragraph, bookmark_id):
        """Adiciona hyperlink interno a um parágrafo apontando para bookmark"""
        # Salvar o texto original
        original_text = paragraph.text

        # Limpar TODO o conteúdo do parágrafo (runs, hyperlinks, etc)
        for child in list(paragraph._element):
            if child.tag in [qn('w:r'), qn('w:hyperlink')]:
                paragraph._element.remove(child)

        # Criar elemento de hyperlink
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('w:anchor'), bookmark_id)

        # Criar um novo run com o texto
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        # Estilo de hyperlink (azul, sublinhado)
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0563C1')  # Azul padrão de hyperlink
        rPr.append(color)

        underline = OxmlElement('w:u')
        underline.set(qn('w:val'), 'single')
        rPr.append(underline)

        new_run.append(rPr)

        # Adicionar texto
        text_elem = OxmlElement('w:t')
        text_elem.text = original_text
        new_run.append(text_elem)

        # Montar hyperlink
        hyperlink.append(new_run)

        # Adicionar ao parágrafo
        paragraph._element.append(hyperlink)

    def _add_bookmark_to_paragraph(self, paragraph, bookmark_id):
        """Adiciona um bookmark a um parágrafo (se não existir)"""
        # Verificar se já existe bookmark
        bookmarks = paragraph._element.findall('.//' + qn('w:bookmarkStart'))
        for bm in bookmarks:
            if bm.get(qn('w:name')) == bookmark_id:
                return  # Já existe

        # Gerar ID único para o bookmark
        import random
        bookmark_internal_id = str(random.randint(1000000, 9999999))

        # Criar bookmarkStart
        bookmark_start = OxmlElement('w:bookmarkStart')
        bookmark_start.set(qn('w:id'), bookmark_internal_id)
        bookmark_start.set(qn('w:name'), bookmark_id)

        # Criar bookmarkEnd
        bookmark_end = OxmlElement('w:bookmarkEnd')
        bookmark_end.set(qn('w:id'), bookmark_internal_id)

        # Adicionar ao início e fim do parágrafo
        paragraph._element.insert(0, bookmark_start)
        paragraph._element.append(bookmark_end)

    def customize_docx(self, docx_path: Path):
        """Customiza estilos do DOCX após conversão"""
        doc = Document(str(docx_path))

        # Configurar margens (0.5 polegadas em todos os lados)
        self.set_margins(doc, left=0.5, right=0.5, top=0.5, bottom=0.5)

        # Customizar estilos de fonte
        styles = doc.styles

        # Estilo Normal (texto padrão)
        if 'Normal' in styles:
            style = styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(11)

        # Estilo Title (título do documento) - Arial
        if 'Title' in styles:
            style = styles['Title']
            style.font.name = 'Arial'
            style.font.size = Pt(18)

        # Estilos de Heading (títulos) - todos em Arial
        for i in range(1, 10):
            heading_name = f'Heading {i}'
            if heading_name in styles:
                style = styles[heading_name]
                style.font.name = 'Arial'
                # Tamanhos decrescentes para headings
                sizes = {1: 16, 2: 14, 3: 13, 4: 12, 5: 11, 6: 11, 7: 11, 8: 11, 9: 11}
                style.font.size = Pt(sizes.get(i, 11))

        # Estilos de código inline (entre ` `) - Courier New
        inline_code_styles = ['Verbatim Char']
        for style_name in inline_code_styles:
            if style_name in styles:
                style = styles[style_name]
                style.font.name = 'Courier New'
                style.font.size = Pt(11)

        # Estilos de blocos de código - Courier New (fonte monoespaçada)
        # NOTA: NÃO incluir 'Compact' aqui (Pandoc usa para listas, não código)
        block_code_styles = ['Source Code', 'Verbatim', 'Code',
                             'Preformatted Text', 'HTML Preformatted']
        for style_name in block_code_styles:
            if style_name in styles:
                style = styles[style_name]
                style.font.name = 'Courier New'
                style.font.size = Pt(11)

        # Estilos de lista - Arial (Pandoc usa 'Compact' para listas)
        list_styles = ['Compact', 'List Paragraph', 'List', 'List Bullet', 'List Number']
        for style_name in list_styles:
            if style_name in styles:
                style = styles[style_name]
                style.font.name = 'Arial'
                style.font.size = Pt(11)

        # Customizar tabelas (apenas bordas horizontais)
        self.customize_tables(doc)

        # Garantir formatação de listas (bullets e números) - ANTES de aplicar fundo
        self.ensure_list_formatting(doc)

        # Adicionar fundo cinza APENAS aos blocos de código (não listas)
        self.customize_code_blocks(doc)

        # Corrigir TOC: remover título "Sumário" e adicionar hyperlinks
        self.fix_toc(doc)

        # Salvar customizações
        doc.save(str(docx_path))
        print(f"  [OK] Estilos customizados aplicados (fontes + tabelas + codigo + listas + TOC)")

    def convert(
        self,
        md_content: str,
        output_path: Path,
        title: str = None,
        author: str = "DestaquesGovBr Team"
    ):
        """Converte Markdown para DOCX"""
        # Cria arquivo temporário
        with tempfile.NamedTemporaryFile(
            mode='w',
            suffix='.md',
            delete=False,
            encoding='utf-8'
        ) as tmp:
            tmp.write(md_content)
            tmp_path = tmp.name

        try:
            # Caminho do arquivo de referência (template DOCX com fontes customizadas)
            script_dir = Path(__file__).parent
            reference_docx = script_dir / 'templates' / 'reference.docx'

            # Metadata para Pandoc
            extra_args = [
                '--standalone',
                '--wrap=auto',
                '--highlight-style=pygments',  # Syntax highlighting com cores
            ]

            # Usar arquivo de referência se existir
            if reference_docx.exists():
                extra_args.append(f'--reference-doc={reference_docx}')

            if title:
                extra_args.extend(['-M', f'title={title}'])

            extra_args.extend(['-M', f'author={author}'])

            # Conversão via pypandoc
            pypandoc.convert_file(
                tmp_path,
                'docx',
                outputfile=str(output_path),
                extra_args=extra_args
            )

            # Customizar estilos
            self.customize_docx(output_path)

            print(f"  [OK] DOCX gerado: {output_path}")

        finally:
            Path(tmp_path).unlink(missing_ok=True)


def convert_single_file(
    input_path: Path,
    output_path: Path = None,
    cache_dir: Path = None
):
    """Converte um único arquivo Markdown para DOCX"""

    # Paths padrão
    if output_path is None:
        output_path = input_path.parent / OUTPUT_DIR / (
            input_path.stem + '.docx'
        )

    if cache_dir is None:
        cache_dir = input_path.parent / OUTPUT_DIR / IMG_DIR

    # Cria diretórios de output
    output_path.parent.mkdir(parents=True, exist_ok=True)

    print(f"\n{'='*60}")
    print(f"Convertendo: {input_path.name}")
    print(f"{'='*60}")

    # Lê o arquivo Markdown
    md_content = input_path.read_text(encoding='utf-8')

    # Processa Mermaid
    renderer = MermaidRenderer(cache_dir)
    processor = MarkdownProcessor(renderer)

    # Path ABSOLUTO das imagens (para que Pandoc encontre de qualquer lugar)
    img_path = str(cache_dir.absolute()).replace('\\', '/')
    processed_md = processor.process(md_content, img_path)

    # Converte para DOCX
    converter = DOCXConverter()
    title = input_path.stem.replace('-', ' ').title()
    converter.convert(processed_md, output_path, title=title)

    print(f"\n[OK] Conversao concluida!")
    print(f"   Output: {output_path}")


def convert_all_files(relatorios_dir: Path):
    """Converte todos os arquivos .md no diretório"""
    md_files = sorted(list(relatorios_dir.glob('*.md')))

    if not md_files:
        print(f"[X] Nenhum arquivo .md encontrado em: {relatorios_dir}")
        return 1

    print(f"\nEncontrados {len(md_files)} arquivos Markdown:")
    for f in md_files:
        print(f"  - {f.name}")

    success_count = 0
    error_count = 0

    for md_file in md_files:
        try:
            convert_single_file(md_file)
            success_count += 1
        except Exception as e:
            print(f"\n[X] Erro ao converter {md_file.name}: {e}")
            error_count += 1
            continue

    print(f"\n{'='*60}")
    print(f"Resumo: {success_count} sucesso(s), {error_count} erro(s)")
    print(f"{'='*60}")

    return 0 if error_count == 0 else 1


def main():
    parser = argparse.ArgumentParser(
        description='Converte relatórios Markdown para DOCX com diagramas Mermaid renderizados',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Exemplos:
  %(prog)s relatorio.md
  %(prog)s relatorio.md --output meu-relatorio.docx
  %(prog)s --all
  %(prog)s --all --relatorios-dir docs/docs/relatorios/
        """
    )
    parser.add_argument(
        'input',
        nargs='?',
        type=Path,
        help='Arquivo .md de entrada'
    )
    parser.add_argument(
        '--output', '-o',
        type=Path,
        help='Arquivo .docx de saída (opcional)'
    )
    parser.add_argument(
        '--all',
        action='store_true',
        help='Converte todos os .md no diretório de relatórios'
    )
    parser.add_argument(
        '--relatorios-dir',
        type=Path,
        default=Path('docs/docs/relatorios'),
        help='Diretório contendo os relatórios (padrão: docs/docs/relatorios)'
    )

    args = parser.parse_args()

    # Verifica dependências do sistema
    check_dependencies()

    # Validação
    if args.all:
        relatorios_dir = args.relatorios_dir
        if not relatorios_dir.exists():
            print(f"[X] Diretório não encontrado: {relatorios_dir}")
            return 1

        return convert_all_files(relatorios_dir)

    elif args.input:
        if not args.input.exists():
            print(f"[X] Arquivo não encontrado: {args.input}")
            return 1

        convert_single_file(args.input, args.output)
        return 0

    else:
        parser.print_help()
        return 1


if __name__ == '__main__':
    sys.exit(main())
