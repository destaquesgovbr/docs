#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Merge COMPLETO de DOCX usando docxcompose
Preserva TODOS os estilos, imagens, fontes, tabelas e configurações

Uso:
    python merge_docx_with_docxcompose.py <template.docx> <content.docx> <output.docx>
"""

import sys
from pathlib import Path

try:
    from docx import Document
    from docxcompose.composer import Composer
except ImportError as e:
    print("[ERRO] Bibliotecas nao instaladas.")
    print("   Instale com: pip install python-docx docxcompose")
    print(f"   Erro: {e}")
    sys.exit(2)


def merge_docx_with_composer(template_path, content_path, output_path):
    """
    Faz merge usando docxcompose que preserva TUDO automaticamente
    """
    print(f"\n{'='*60}")
    print(f"MERGE COM DOCXCOMPOSE")
    print(f"{'='*60}")
    print(f"Template: {template_path}")
    print(f"Conteudo: {content_path}")
    print(f"Output: {output_path}")
    print(f"{'='*60}\n")

    print(f"  [OK] Carregando template...")
    doc_template = Document(template_path)

    print(f"  [OK] Carregando conteudo...")
    doc_content = Document(content_path)

    # Estatísticas antes do merge
    print(f"\n  [INFO] Template:")
    print(f"    - Paragrafos: {len(doc_template.paragraphs)}")
    print(f"    - Tabelas: {len(doc_template.tables)}")
    print(f"    - Secoes: {len(doc_template.sections)}")

    print(f"\n  [INFO] Conteudo:")
    print(f"    - Paragrafos: {len(doc_content.paragraphs)}")
    print(f"    - Tabelas: {len(doc_content.tables)}")
    print(f"    - Secoes: {len(doc_content.sections)}")

    # Usa Composer para fazer merge preservando tudo
    print(f"\n  [OK] Fazendo merge com docxcompose...")
    print(f"      (preserva estilos, imagens, fontes, configuracoes)")

    # Cria composer a partir do template
    composer = Composer(doc_template)

    # Adiciona o documento de conteúdo
    composer.append(doc_content)

    print(f"  [OK] Merge concluido")

    # Salva resultado
    print(f"\n  [OK] Salvando documento mesclado...")
    composer.save(output_path)

    # Estatísticas finais
    print(f"\n  [INFO] Documento final:")
    doc_final = Document(output_path)
    print(f"    - Paragrafos: {len(doc_final.paragraphs)}")
    print(f"    - Tabelas: {len(doc_final.tables)}")
    print(f"    - Secoes: {len(doc_final.sections)}")

    # Contar imagens
    img_count = 0
    for paragraph in doc_final.paragraphs:
        for run in paragraph.runs:
            if run._element.xpath('.//pic:pic'):
                img_count += 1

    # Imagens em tabelas
    for table in doc_final.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run._element.xpath('.//pic:pic'):
                            img_count += 1

    print(f"    - Imagens: {img_count}")

    # Verificar arquivos de mídia no ZIP
    import zipfile
    with zipfile.ZipFile(output_path, 'r') as z:
        media_files = [f for f in z.namelist() if 'media/' in f]
        print(f"    - Arquivos de midia no ZIP: {len(media_files)}")

    print(f"\n{'='*60}")
    print(f"[OK] MERGE COMPLETO COM DOCXCOMPOSE CONCLUIDO!")
    print(f"{'='*60}\n")
    print(f"Arquivo gerado: {output_path}")
    print(f"\n[INFO] docxcompose preserva automaticamente:")
    print(f"  - Todos os estilos e fontes")
    print(f"  - Todas as imagens (incluindo do header/footer)")
    print(f"  - Todas as tabelas com formatacao")
    print(f"  - Configuracoes de pagina")
    print(f"  - Numeracao de listas")
    print(f"  - Relationships (links entre elementos e imagens)")


def main():
    if len(sys.argv) != 4:
        print("Uso: python merge_docx_with_docxcompose.py <template.docx> <content.docx> <output.docx>")
        print("\nExemplo:")
        print('  python merge_docx_with_docxcompose.py template.docx relatorio.docx resultado.docx')
        sys.exit(1)

    template_path = sys.argv[1]
    content_path = sys.argv[2]
    output_path = sys.argv[3]

    # Valida arquivos de entrada
    if not Path(template_path).exists():
        print(f"[ERRO] Template nao encontrado: {template_path}")
        sys.exit(1)

    if not Path(content_path).exists():
        print(f"[ERRO] Arquivo de conteudo nao encontrado: {content_path}")
        sys.exit(1)

    # Cria diretório de saída se não existir
    output_dir = Path(output_path).parent
    output_dir.mkdir(parents=True, exist_ok=True)

    try:
        merge_docx_with_composer(template_path, content_path, output_path)
    except Exception as e:
        print(f"\n[ERRO] Falha ao fazer merge: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()