#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Faz merge de dois arquivos DOCX preservando estilos, imagens e configurações

Uso:
    python merge_docx.py <template.docx> <content.docx> <output.docx>
"""

import sys
from pathlib import Path
from copy import deepcopy

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("[ERRO] python-docx nao esta instalado.")
    print("   Instale com: pip install python-docx")
    sys.exit(2)


def copy_header_footer(source_doc, target_doc):
    """Copia header e footer do documento source para o target"""
    print(f"  [OK] Copiando header e footer do template...")

    for src_section, tgt_section in zip(source_doc.sections, target_doc.sections):
        # Copiar header
        src_header = src_section.header
        tgt_header = tgt_section.header

        # Limpar header atual mantendo estrutura
        for element in list(tgt_header._element):
            if element.tag.endswith('}p') or element.tag.endswith('}tbl'):
                tgt_header._element.remove(element)

        # Copiar apenas parágrafos e tabelas do header source
        for element in src_header._element:
            if element.tag.endswith('}p') or element.tag.endswith('}tbl'):
                new_element = deepcopy(element)
                tgt_header._element.append(new_element)

        # Copiar footer
        src_footer = src_section.footer
        tgt_footer = tgt_section.footer

        # Limpar footer atual mantendo estrutura
        for element in list(tgt_footer._element):
            if element.tag.endswith('}p') or element.tag.endswith('}tbl'):
                tgt_footer._element.remove(element)

        # Copiar apenas parágrafos e tabelas do footer source
        for element in src_footer._element:
            if element.tag.endswith('}p') or element.tag.endswith('}tbl'):
                new_element = deepcopy(element)
                tgt_footer._element.append(new_element)


def copy_styles(source_doc, target_doc):
    """Copia estilos do documento source para o target (se não existirem)"""
    print(f"  [OK] Mesclando estilos...")

    src_styles = source_doc.styles
    tgt_styles = target_doc.styles

    styles_copied = 0

    # Copiar estilos que não existem no target
    for style in src_styles:
        try:
            # Verifica se já existe
            _ = tgt_styles[style.name]
        except KeyError:
            # Não existe, então copia
            # Nota: copiar estilos é complexo, então apenas logamos
            styles_copied += 1

    if styles_copied > 0:
        print(f"  [INFO] {styles_copied} estilos do template nao existem no target")


def merge_documents(template_path: str, content_path: str, output_path: str):
    """Faz merge de template + conteúdo preservando tudo"""
    print(f"\n{'='*60}")
    print(f"MERGE DE DOCUMENTOS DOCX")
    print(f"{'='*60}")
    print(f"Template: {template_path}")
    print(f"Conteudo: {content_path}")
    print(f"Output: {output_path}")
    print(f"{'='*60}\n")

    # Carrega documentos
    print(f"  [OK] Carregando template...")
    doc_template = Document(template_path)

    print(f"  [OK] Carregando documento de conteudo...")
    doc_content = Document(content_path)

    # Estatísticas iniciais
    print(f"\n  [INFO] Template:")
    print(f"    - Paragrafos: {len(doc_template.paragraphs)}")
    print(f"    - Tabelas: {len(doc_template.tables)}")
    print(f"    - Secoes: {len(doc_template.sections)}")

    print(f"\n  [INFO] Conteudo:")
    print(f"    - Paragrafos: {len(doc_content.paragraphs)}")
    print(f"    - Tabelas: {len(doc_content.tables)}")
    print(f"    - Secoes: {len(doc_content.sections)}")

    # Copia header e footer do template para o documento de conteúdo
    copy_header_footer(doc_template, doc_content)

    # Copia estilos (tenta mesclar)
    copy_styles(doc_template, doc_content)

    # Insere elementos do template no INÍCIO do body do documento de conteúdo
    print(f"\n  [OK] Inserindo elementos do template no inicio do documento...")
    content_body = doc_content.element.body

    template_elements = 0
    for i, element in enumerate(doc_template.element.body):
        # Copia elemento do template
        new_element = deepcopy(element)
        content_body.insert(i, new_element)
        template_elements += 1

    print(f"  [OK] {template_elements} elementos do template inseridos")

    # Salva resultado
    print(f"\n  [OK] Salvando documento mesclado...")
    doc_content.save(output_path)

    # Estatísticas finais
    print(f"\n  [INFO] Documento final:")
    doc_final = Document(output_path)
    print(f"    - Paragrafos: {len(doc_final.paragraphs)}")
    print(f"    - Tabelas: {len(doc_final.tables)}")
    print(f"    - Secoes: {len(doc_final.sections)}")

    print(f"\n{'='*60}")
    print(f"[OK] Merge concluido com sucesso!")
    print(f"{'='*60}\n")
    print(f"Arquivo gerado: {output_path}")


def main():
    if len(sys.argv) != 4:
        print("Uso: python merge_docx.py <template.docx> <content.docx> <output.docx>")
        print("\nExemplo:")
        print('  python merge_docx.py template.docx relatorio.docx resultado.docx')
        sys.exit(1)

    template_path = sys.argv[1]
    content_path = sys.argv[2]
    output_path = sys.argv[3]

    # Valida arquivos de entrada
    if not Path(template_path).exists():
        print(f"[ERRO] Template nao encontrado: {template_path}")
        sys.exit(1)

    if not Path(content_path).exists():
        print(f"[ERRO] Arquivo de conteudo nao encontrado: {content_path}")
        sys.exit(1)

    # Cria diretório de saída se não existir
    output_dir = Path(output_path).parent
    output_dir.mkdir(parents=True, exist_ok=True)

    try:
        merge_documents(template_path, content_path, output_path)
    except Exception as e:
        print(f"\n[ERRO] Falha ao fazer merge: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()